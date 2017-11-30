# -*- coding: utf-8 -*-
#
# 2017.11.22：在addText中增加color参数
#

from docx import Document
from docx.shared import Pt
from docx.shared import Inches
from docx.shared import RGBColor
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

import sys

reload(sys)
sys.setdefaultencoding('utf-8')

class createWord:

    def __init__(self):

        self.document = Document()
        self.document.styles['Normal'].font.name = u'微软黑体'
        self.document.styles['Normal'].paragraph_format.space_before = Pt(10)
        self.document.styles['Normal'].paragraph_format.space_after = Pt(10)
        self.document.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'微软黑体')
        self.document.styles['Title'].font.name = u'微软黑体'
        self.document.styles['Heading1'].font.name = u'微软黑体'
        self.document.styles['Heading2'].font.name = u'微软黑体'
        self.document.styles['Heading3'].font.name = u'微软黑体'

    """Paragraph：段落"""
    def addHead(self, info, lvl, align=WD_ALIGN_PARAGRAPH.LEFT):
        _style = 'Title' if lvl == 0 else 'Heading%d' % lvl
        _paragraph = self.document.add_paragraph(info, _style)
        _paragraph.paragraph_format.alignment = align

    def addText(self, info, align=WD_ALIGN_PARAGRAPH.LEFT, color=None):
        self.paragrap = self.document.add_paragraph()
        _run = self.paragrap.add_run()
        self.paragrap.paragraph_format.alignment = align
        self.paragrap.paragraph_format.first_line_indent = Inches(0.3)
        if color is not None:
            font = _run.font
            font.color.rgb = RGBColor(color[0],color[1],color[2])
        _run.add_text(info)

    def addPic(self, pic_path, sizeof=5):
        self.paragrap = self.document.add_paragraph()
        self.paragrap.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _run = self.paragrap.add_run()
        _run.add_picture(pic_path, width=Inches(sizeof))

    def addTable(self, n_rows, n_cols, col_width=None):
        self.table = self.document.add_table(rows=n_rows, cols=n_cols)
        self.table.alignment = WD_TABLE_ALIGNMENT.CENTER
        self.table.style = 'MediumList1'
        self.rows = n_rows
        self.cols = n_cols
        self._idx = 0
        if col_width is not None:
            for _i in range(n_cols):
                self.table.cell(0,_i).width = col_width[_i] * 914400

    def addRow(self,data):
        if self._idx < self.rows:
            _hdr_cells = self.table.rows[self._idx].cells
            for __idx in range(self.cols):
                _run = _hdr_cells[__idx].paragraphs[0].add_run()
                if data[__idx][0] == 'pic':
                    _run.add_picture(data[__idx][1], width=Inches(1.2))
                else:
                    _run.font.size = Pt(10)
                    _hdr_cells[__idx].text = data[__idx][1]
        else:
            _hdr_cells = self.table.add_row().cells
            for __idx in range(self.cols):
                _run = _hdr_cells[__idx].paragraphs[0].add_run()
                if data[__idx][0] == 'pic':
                    _run.add_picture(data[__idx][1], width=Inches(data[__idx][2]))
                else:
                    _hdr_cells[__idx].text = data[__idx][1]
        self._idx += 1

    def setTableFont(self,ft):
        for row in self.table.rows:
            for cell in row.cells:
                paragraphs = cell.paragraphs
                for paragraph in paragraphs:
                    paragraph.paragraph_format.space_before = Pt(2)
                    paragraph.paragraph_format.space_after = Pt(2)
                    for run in paragraph.runs:
                        font = run.font
                        font.size = Pt(ft)

    def addPageBreak(self):
        self.document.add_page_break()

    def saveFile(self, fname):
        self.document.save(fname)

#
#